"""
Document Editor Module for AI Memory Storage System

Provides editing capabilities for various document formats:
- Microsoft Word (.docx)
- Microsoft Excel (.xlsx)
- Apple Pages (.pages)
- Apple Numbers (.numbers)
- Text files

Each editor implements the following interface:
- load(file_path): Load document from file_path
- get_content(): Get document content as text or structured data
- edit_content(edit_instructions): Edit document based on instructions
- save(output_path): Save document to output_path
"""

import os
import io
import re
import tempfile
import zipfile
import shutil
import logging
import subprocess
from typing import Dict, Any, List, Optional, Union
import json

# Setup logging
logger = logging.getLogger(__name__)

# Document editors for different file types
try:
    import docx
    from docx.shared import Pt
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    
try:
    import openpyxl
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False

try:
    from lxml import etree
    XML_SUPPORT = True
except ImportError:
    XML_SUPPORT = False

# Base Editor class
class DocumentEditor:
    """Base class for document editors"""
    
    def __init__(self):
        self.content = None
        self.file_path = None
        self.metadata = {}
        
    def load(self, file_path: str) -> bool:
        """Load document from file path"""
        self.file_path = file_path
        return os.path.exists(file_path)
        
    def get_content(self) -> str:
        """Get document content"""
        return ""
        
    def edit_content(self, edit_instructions: Dict[str, Any]) -> bool:
        """Edit document based on instructions"""
        return False
        
    def save(self, output_path: str) -> bool:
        """Save document to output path"""
        return False
        
    def get_metadata(self) -> Dict[str, Any]:
        """Get document metadata"""
        return self.metadata

# Plain Text Editor
class TextEditor(DocumentEditor):
    """Editor for plain text files"""
    
    def load(self, file_path: str) -> bool:
        """Load text file"""
        super().load(file_path)
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                self.content = f.read()
            return True
        except Exception as e:
            logger.error(f"Error loading text file: {e}")
            return False
            
    def get_content(self) -> str:
        """Get text content"""
        return self.content or ""
        
    def edit_content(self, edit_instructions: Dict[str, Any]) -> bool:
        """Edit text content based on instructions"""
        try:
            # Handle different edit operations
            op = edit_instructions.get('operation', 'replace_all')
            
            if op == 'replace_all':
                # Complete replacement
                new_content = edit_instructions.get('content', '')
                if new_content:
                    self.content = new_content
                    return True
                    
            elif op == 'find_replace':
                # Find and replace text
                find_text = edit_instructions.get('find', '')
                replace_text = edit_instructions.get('replace', '')
                if find_text and self.content:
                    self.content = self.content.replace(find_text, replace_text)
                    return True
                    
            elif op == 'regex_replace':
                # Regex find and replace
                pattern = edit_instructions.get('pattern', '')
                replace_text = edit_instructions.get('replace', '')
                if pattern and self.content:
                    self.content = re.sub(pattern, replace_text, self.content)
                    return True
                    
            elif op == 'append':
                # Append text
                append_text = edit_instructions.get('content', '')
                if append_text:
                    self.content = (self.content or "") + append_text
                    return True
                    
            elif op == 'insert_at':
                # Insert at position
                insert_text = edit_instructions.get('content', '')
                position = edit_instructions.get('position', 0)
                if insert_text and isinstance(position, int):
                    pos = max(0, min(position, len(self.content or "")))
                    self.content = (self.content or "")[:pos] + insert_text + (self.content or "")[pos:]
                    return True
            
            return False
        except Exception as e:
            logger.error(f"Error editing text content: {e}")
            return False
            
    def save(self, output_path: str) -> bool:
        """Save text file"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(self.content or "")
            return True
        except Exception as e:
            logger.error(f"Error saving text file: {e}")
            return False

# Word Document Editor
class DocxEditor(DocumentEditor):
    """Editor for Microsoft Word .docx files"""
    
    def __init__(self):
        super().__init__()
        self.doc = None
        
    def load(self, file_path: str) -> bool:
        """Load Word document"""
        if not DOCX_SUPPORT:
            logger.error("python-docx not installed")
            return False
            
        super().load(file_path)
        try:
            self.doc = docx.Document(file_path)
            return True
        except Exception as e:
            logger.error(f"Error loading Word document: {e}")
            return False
            
    def get_content(self) -> str:
        """Get document content as text"""
        if not self.doc:
            return ""
            
        try:
            paragraphs = [p.text for p in self.doc.paragraphs]
            tables_text = []
            
            # Extract text from tables
            for table in self.doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    tables_text.append(" | ".join(row_text))
            
            # Combine text
            all_text = "\n".join(paragraphs + tables_text)
            return all_text
        except Exception as e:
            logger.error(f"Error getting Word document content: {e}")
            return ""
            
    def edit_content(self, edit_instructions: Dict[str, Any]) -> bool:
        """Edit Word document based on instructions"""
        if not self.doc:
            return False
            
        try:
            op = edit_instructions.get('operation', '')
            
            if op == 'replace_all_text':
                # Replace all text with new content
                new_content = edit_instructions.get('content', '')
                if new_content:
                    # Clear existing paragraphs
                    for i in range(len(self.doc.paragraphs) - 1, -1, -1):
                        p = self.doc.paragraphs[i]
                        p.clear()
                    
                    # Add new paragraphs
                    if len(self.doc.paragraphs) > 0:
                        # Use the first paragraph
                        self.doc.paragraphs[0].add_run(new_content)
                    else:
                        # Create a new paragraph
                        self.doc.add_paragraph(new_content)
                    return True
                    
            elif op == 'find_replace':
                # Find and replace text
                find_text = edit_instructions.get('find', '')
                replace_text = edit_instructions.get('replace', '')
                if find_text:
                    # Replace in paragraphs
                    for p in self.doc.paragraphs:
                        if find_text in p.text:
                            inline = p.runs
                            # Replace text in each run
                            for i in range(len(inline)):
                                if find_text in inline[i].text:
                                    inline[i].text = inline[i].text.replace(find_text, replace_text)
                    
                    # Replace in tables
                    for table in self.doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    if find_text in p.text:
                                        inline = p.runs
                                        for i in range(len(inline)):
                                            if find_text in inline[i].text:
                                                inline[i].text = inline[i].text.replace(find_text, replace_text)
                    return True
                    
            elif op == 'add_paragraph':
                # Add a new paragraph
                content = edit_instructions.get('content', '')
                position = edit_instructions.get('position', -1)  # -1 means append at end
                
                if content:
                    if position < 0 or position >= len(self.doc.paragraphs):
                        # Append at end
                        self.doc.add_paragraph(content)
                    else:
                        # Insert at position - more complex, would need to rebuild document
                        # This is a simplified approach
                        p = self.doc.paragraphs[position]
                        p.insert_paragraph_before(content)
                    return True
            
            return False
        except Exception as e:
            logger.error(f"Error editing Word document: {e}")
            return False
            
    def save(self, output_path: str) -> bool:
        """Save Word document"""
        if not self.doc:
            return False
            
        try:
            self.doc.save(output_path)
            return True
        except Exception as e:
            logger.error(f"Error saving Word document: {e}")
            return False

# Excel Document Editor
class ExcelEditor(DocumentEditor):
    """Editor for Microsoft Excel .xlsx files"""
    
    def __init__(self):
        super().__init__()
        self.workbook = None
        
    def load(self, file_path: str) -> bool:
        """Load Excel document"""
        if not EXCEL_SUPPORT:
            logger.error("openpyxl not installed")
            return False
            
        super().load(file_path)
        try:
            self.workbook = openpyxl.load_workbook(file_path)
            return True
        except Exception as e:
            logger.error(f"Error loading Excel document: {e}")
            return False
            
    def get_content(self) -> Dict[str, List[List[Any]]]:
        """Get Excel content as structured data"""
        if not self.workbook:
            return {}
            
        try:
            content = {}
            for sheet_name in self.workbook.sheetnames:
                sheet = self.workbook[sheet_name]
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    sheet_data.append(list(row))
                content[sheet_name] = sheet_data
            return content
        except Exception as e:
            logger.error(f"Error getting Excel content: {e}")
            return {}
            
    def edit_content(self, edit_instructions: Dict[str, Any]) -> bool:
        """Edit Excel document based on instructions"""
        if not self.workbook:
            return False
            
        try:
            op = edit_instructions.get('operation', '')
            
            if op == 'update_cell':
                # Update a specific cell
                sheet_name = edit_instructions.get('sheet', self.workbook.active.title)
                row = edit_instructions.get('row', 1)
                column = edit_instructions.get('column', 1)
                value = edit_instructions.get('value', '')
                
                if sheet_name in self.workbook:
                    sheet = self.workbook[sheet_name]
                    sheet.cell(row=row, column=column, value=value)
                    return True
                    
            elif op == 'update_range':
                # Update a range of cells
                sheet_name = edit_instructions.get('sheet', self.workbook.active.title)
                start_row = edit_instructions.get('start_row', 1)
                start_col = edit_instructions.get('start_col', 1)
                data = edit_instructions.get('data', [])
                
                if sheet_name in self.workbook and data:
                    sheet = self.workbook[sheet_name]
                    for i, row_data in enumerate(data):
                        for j, value in enumerate(row_data):
                            sheet.cell(row=start_row+i, column=start_col+j, value=value)
                    return True
                    
            elif op == 'add_sheet':
                # Add a new sheet
                sheet_name = edit_instructions.get('name', f'Sheet{len(self.workbook.sheetnames)+1}')
                data = edit_instructions.get('data', [])
                
                if sheet_name not in self.workbook:
                    sheet = self.workbook.create_sheet(sheet_name)
                    for i, row_data in enumerate(data):
                        for j, value in enumerate(row_data):
                            sheet.cell(row=i+1, column=j+1, value=value)
                    return True
                    
            elif op == 'delete_sheet':
                # Delete a sheet
                sheet_name = edit_instructions.get('name', '')
                
                if sheet_name in self.workbook and len(self.workbook.sheetnames) > 1:
                    del self.workbook[sheet_name]
                    return True
            
            return False
        except Exception as e:
            logger.error(f"Error editing Excel document: {e}")
            return False
            
    def save(self, output_path: str) -> bool:
        """Save Excel document"""
        if not self.workbook:
            return False
            
        try:
            self.workbook.save(output_path)
            return True
        except Exception as e:
            logger.error(f"Error saving Excel document: {e}")
            return False

# Apple Pages Editor
class PagesEditor(DocumentEditor):
    """Editor for Apple Pages documents"""
    
    def __init__(self):
        super().__init__()
        self.temp_dir = None
        self.is_package = False
        self.content_text = ""
        self.has_textutil = False
        
    def __del__(self):
        """Clean up temporary directory"""
        self._cleanup()
        
    def _cleanup(self):
        """Clean up temporary directory"""
        if self.temp_dir and os.path.exists(self.temp_dir):
            try:
                shutil.rmtree(self.temp_dir, ignore_errors=True)
                self.temp_dir = None
            except:
                pass
        
    def load(self, file_path: str) -> bool:
        """Load Pages document"""
        super().load(file_path)
        self._cleanup()
        
        try:
            # Check if textutil is available (macOS)
            try:
                result = subprocess.run(['which', 'textutil'], capture_output=True, text=True)
                self.has_textutil = result.returncode == 0
            except:
                self.has_textutil = False
                
            # Create temp directory
            self.temp_dir = tempfile.mkdtemp()
            
            # Extract content using textutil if available
            if self.has_textutil:
                try:
                    result = subprocess.run(['textutil', '-convert', 'txt', '-stdout', file_path], 
                                          capture_output=True, text=True)
                    if result.returncode == 0:
                        self.content_text = result.stdout
                        self.is_package = True
                        return True
                except:
                    pass
                    
            # Try to extract as a zip package
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    zip_ref.extractall(self.temp_dir)
                    self.is_package = True
                    
                    # Try to find content in index.xml
                    index_path = os.path.join(self.temp_dir, 'index.xml')
                    if os.path.exists(index_path) and XML_SUPPORT:
                        tree = etree.parse(index_path)
                        text_nodes = tree.xpath('//text')
                        self.content_text = '\n'.join([node.text for node in text_nodes if node.text])
                        return True
                        
                    # Try Preview.pdf as fallback
                    preview_path = os.path.join(self.temp_dir, 'QuickLook', 'Preview.pdf')
                    if os.path.exists(preview_path) and PDF_SUPPORT:
                        from PyPDF2 import PdfReader
                        with open(preview_path, 'rb') as f:
                            pdf = PdfReader(f)
                            self.content_text = ""
                            for page in pdf.pages:
                                self.content_text += page.extract_text() + "\n"
                            return True
                            
                    return False
            except:
                # Not a zip file or package
                return False
                
        except Exception as e:
            logger.error(f"Error loading Pages document: {e}")
            return False
            
    def get_content(self) -> str:
        """Get document content as text"""
        return self.content_text
            
    def edit_content(self, edit_instructions: Dict[str, Any]) -> bool:
        """Edit Pages document based on instructions"""
        # Pages documents are complex packages - direct editing is limited
        # We can extract text, modify it, and create a new text file with the edited content
        # But reassembling a proper Pages document is beyond this scope
        
        if not self.content_text:
            return False
            
        try:
            op = edit_instructions.get('operation', '')
            
            if op == 'export_as_text':
                # We already have the text content
                return True
                
            # Other operations aren't supported directly
            logger.warning("Direct editing of Pages documents is not supported. Use export_as_text operation instead.")
            return False
            
        except Exception as e:
            logger.error(f"Error editing Pages document: {e}")
            return False
            
    def save(self, output_path: str) -> bool:
        """Save Pages document (as text)"""
        try:
            # We can only save as text
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(self.content_text or "")
            return True
        except Exception as e:
            logger.error(f"Error saving Pages document as text: {e}")
            return False

# Apple Numbers Editor
class NumbersEditor(DocumentEditor):
    """Editor for Apple Numbers documents"""
    
    def __init__(self):
        super().__init__()
        self.temp_dir = None
        self.content_data = {}
        self.has_textutil = False
        
    def __del__(self):
        """Clean up temporary directory"""
        self._cleanup()
        
    def _cleanup(self):
        """Clean up temporary directory"""
        if self.temp_dir and os.path.exists(self.temp_dir):
            try:
                shutil.rmtree(self.temp_dir, ignore_errors=True)
                self.temp_dir = None
            except:
                pass
        
    def load(self, file_path: str) -> bool:
        """Load Numbers document"""
        super().load(file_path)
        self._cleanup()
        
        try:
            # Check if textutil is available (macOS)
            try:
                result = subprocess.run(['which', 'textutil'], capture_output=True, text=True)
                self.has_textutil = result.returncode == 0
            except:
                self.has_textutil = False
                
            # Create temp directory
            self.temp_dir = tempfile.mkdtemp()
            
            # Extract content using textutil if available
            if self.has_textutil:
                try:
                    result = subprocess.run(['textutil', '-convert', 'txt', '-stdout', file_path], 
                                          capture_output=True, text=True)
                    if result.returncode == 0:
                        text_content = result.stdout
                        
                        # Try to parse table-like structure from the text
                        sheets = {}
                        current_sheet = None
                        current_data = []
                        
                        for line in text_content.splitlines():
                            if line.strip().endswith(':') and not line.strip().startswith('•'):
                                # This looks like a sheet name
                                if current_sheet and current_data:
                                    sheets[current_sheet] = current_data
                                current_sheet = line.strip().rstrip(':')
                                current_data = []
                            elif line.strip() and current_sheet is not None:
                                # Try to split by tabs or multiple spaces
                                row = re.split(r'\t+|\s{2,}', line.strip())
                                if row:
                                    current_data.append(row)
                                    
                        # Add the last sheet
                        if current_sheet and current_data:
                            sheets[current_sheet] = current_data
                            
                        self.content_data = sheets
                        return True
                except:
                    pass
                    
            # Try to extract as a zip package
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    zip_ref.extractall(self.temp_dir)
                    
                    # Try Preview.pdf as fallback
                    preview_path = os.path.join(self.temp_dir, 'QuickLook', 'Preview.pdf')
                    if os.path.exists(preview_path) and PDF_SUPPORT:
                        from PyPDF2 import PdfReader
                        with open(preview_path, 'rb') as f:
                            pdf = PdfReader(f)
                            text_content = ""
                            for page in pdf.pages:
                                text_content += page.extract_text() + "\n"
                                
                            # Try to parse table-like structure from the text
                            sheets = {"Sheet1": []}
                            for line in text_content.splitlines():
                                if line.strip():
                                    # Try to split by tabs or multiple spaces
                                    row = re.split(r'\t+|\s{2,}', line.strip())
                                    if row:
                                        sheets["Sheet1"].append(row)
                                        
                            self.content_data = sheets
                            return True
                            
                    return False
            except:
                # Not a zip file or package
                return False
                
        except Exception as e:
            logger.error(f"Error loading Numbers document: {e}")
            return False
            
    def get_content(self) -> Dict[str, List[List[Any]]]:
        """Get document content as structured data"""
        return self.content_data
            
    def edit_content(self, edit_instructions: Dict[str, Any]) -> bool:
        """Edit Numbers document based on instructions"""
        # Similar to Pages, direct editing is limited
        if not self.content_data:
            return False
            
        try:
            op = edit_instructions.get('operation', '')
            
            if op == 'export_as_csv':
                # We already have structured data that can be exported as CSV
                return True
                
            # Other operations aren't supported directly
            logger.warning("Direct editing of Numbers documents is not supported. Use export_as_csv operation instead.")
            return False
            
        except Exception as e:
            logger.error(f"Error editing Numbers document: {e}")
            return False
            
    def save(self, output_path: str) -> bool:
        """Save Numbers document (as CSV)"""
        try:
            # We can only save as CSV
            if not output_path.endswith('.csv'):
                output_path += '.csv'
                
            sheet_name = list(self.content_data.keys())[0] if self.content_data else "Sheet1"
            data = self.content_data.get(sheet_name, [])
            
            with open(output_path, 'w', encoding='utf-8') as f:
                for row in data:
                    f.write(','.join([f'"{str(cell).replace(\'"\', \'""\'")}"' for cell in row]) + '\n')
            return True
        except Exception as e:
            logger.error(f"Error saving Numbers document as CSV: {e}")
            return False

# Factory function to get the appropriate editor
def get_editor_for_file(file_path: str) -> Optional[DocumentEditor]:
    """Get appropriate editor for a file based on extension"""
    if not file_path or not os.path.exists(file_path):
        return None
        
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext in ['.txt', '.md', '.json', '.csv', '.log', '.xml', '.html', '.htm']:
        return TextEditor()
    elif file_ext == '.docx' and DOCX_SUPPORT:
        return DocxEditor()
    elif file_ext == '.xlsx' and EXCEL_SUPPORT:
        return ExcelEditor()
    elif file_ext == '.pages':
        return PagesEditor()
    elif file_ext == '.numbers':
        return NumbersEditor()
    else:
        # Default to text editor for unknown types
        return TextEditor()